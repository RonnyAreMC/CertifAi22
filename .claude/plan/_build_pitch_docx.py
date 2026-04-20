"""Genera Certifai_AI_Minutes_Pitch.docx desde el markdown del pitch.

Ejecutar: python .claude/plan/_build_pitch_docx.py
Output: .claude/plan/Certifai_AI_Minutes_Pitch.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm


PRIMARY = RGBColor(0x16, 0x20, 0x54)   # azul UNEMI
ACCENT = RGBColor(0xD4, 0xAF, 0x37)    # dorado
GRAY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1A, 0x1A, 0x1A)


def add_title(doc, text, size=28, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    size = {1: 20, 2: 16, 3: 13}[level]
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = PRIMARY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    return p


def add_bullet(doc, text, bold_first=None):
    """Agrega bullet; si bold_first es str, esa parte inicial va en negrita."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_first:
        r1 = p.add_run(bold_first)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('• • •')
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def add_flow_box(doc, steps):
    """Recuadro con pasos del flujo."""
    table = doc.add_table(rows=len(steps), cols=1)
    table.style = 'Light Shading Accent 1'
    for i, step in enumerate(steps):
        cell = table.rows[i].cells[0]
        cell.text = step
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
    doc.add_paragraph()


def build():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Portada
    doc.add_paragraph()
    add_title(doc, 'Certifai AI Minutes', size=36)
    add_body(doc, 'El evento termina. El contenido empieza.',
             italic=True, size=14, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_divider(doc)

    # SECCIÓN 1
    add_heading(doc, 'El problema que nadie está resolviendo')
    add_body(doc, 'En la educación continua universitaria pasa algo absurdo:')
    add_quote(
        doc,
        'Se organizan eventos con ponentes de talla internacional. Asisten 200 '
        'personas virtualmente. Se graba la sesión. Y todo ese conocimiento se '
        'muere el mismo día que terminó la reunión.'
    )
    add_body(doc, 'Lo que existe hoy en el mejor de los casos:')
    add_bullet(doc, 'Un video de 2 horas que nadie va a volver a ver.')
    add_bullet(doc, 'Un PDF de asistencia que prueba que alguien "estuvo conectado".')
    add_bullet(doc, 'Un certificado que certifica presencia, no aprendizaje.')
    p = doc.add_paragraph()
    r1 = p.add_run('El 90% del valor educativo ')
    r1.font.size = Pt(11)
    r2 = p.add_run('se pierde')
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(' apenas el organizador cierra Zoom o Meet.')
    r3.font.size = Pt(11)
    add_divider(doc)

    # SECCIÓN 2
    add_heading(doc, 'La pregunta radical')
    p = doc.add_paragraph()
    r = p.add_run('¿Y si cada evento virtual generara, automáticamente y sin intervención humana, el siguiente contenido?')
    r.bold = True
    r.font.size = Pt(12)
    add_bullet(doc, ' publicable, navegable, consultable.',
               bold_first='Un resumen ejecutivo')
    add_bullet(doc, ' con los momentos clave en timestamps clickeables.',
               bold_first='Un timeline interactivo')
    add_bullet(doc, ' generadas por IA para quien no pudo asistir.',
               bold_first='10-15 preguntas y respuestas')
    add_bullet(doc, ': frases memorables, citables, con autor y minuto exacto.',
               bold_first='Los highlights')
    add_bullet(doc, ' etiquetados para búsqueda cruzada entre eventos.',
               bold_first='Temas principales')
    p = doc.add_paragraph()
    r = p.add_run('Todo eso. Publicado automático. En minutos. En tu sitio. Para siempre.')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = PRIMARY
    add_body(doc, 'Eso es lo que construimos.')
    add_divider(doc)

    # SECCIÓN 3 — Cómo funciona
    add_heading(doc, 'Cómo funciona en 30 segundos')
    add_flow_box(doc, [
        '1. Admin crea evento en la plataforma',
        '2. Plataforma genera link de Meet automáticamente',
        '3. Ponentes y asistentes entran con ese link',
        '4. Meet graba y transcribe (nativo de Google)',
        '5. Plataforma detecta transcript nuevo en Drive',
        '6. Claude AI procesa el transcript → resumen + Q&A + timeline',
        '7. Página pública del evento se rellena sola',
        '8. ∞',
    ])
    p = doc.add_paragraph()
    r1 = p.add_run('El ponente no hace nada distinto. El admin no sube archivos. El asistente no llena formularios.')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(' Todo es automático.')
    r2.font.size = Pt(11)
    add_divider(doc)

    # SECCIÓN 4 — Por qué es radical
    add_heading(doc, 'Por qué esto es radical')

    add_heading(doc, '1. Invierte el ciclo de vida del evento', level=2)
    add_body(doc, 'Hoy: evento crea asistencia → asistencia genera certificado → fin.')
    p = doc.add_paragraph()
    r1 = p.add_run('Con esta idea: evento crea ')
    r1.font.size = Pt(11)
    r2 = p.add_run('contenido permanente')
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(' → contenido atrae nuevos estudiantes → genera ')
    r3.font.size = Pt(11)
    r4 = p.add_run('más eventos.')
    r4.bold = True
    r4.font.size = Pt(11)
    add_body(doc, 'Es un volante de crecimiento, no una métrica de cumplimiento.')

    add_heading(doc, '2. Hace irrelevante la asistencia como KPI', level=2)
    add_body(
        doc,
        'Para eventos virtuales, "quién se conectó" es una métrica mentirosa. '
        'Mucha gente entra y se va a hacer otra cosa. Lo que importa es: '
        '¿qué se aprendió? Este sistema lo mide en contenido, no en logs de conexión.'
    )

    add_heading(doc, '3. Democratiza el conocimiento internacional', level=2)
    p = doc.add_paragraph()
    r1 = p.add_run('Un ponente del MIT da una charla a 200 personas en Milagro, Ecuador. '
                   'Antes: 200 lo veían. Ahora: ')
    r1.font.size = Pt(11)
    r2 = p.add_run('200 lo viven, 20,000 lo consultan.')
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(' El resumen y Q&A se indexan en Google. Un estudiante de Loja '
                   'que busca "IA ética" encuentra esa charla meses después.')
    r3.font.size = Pt(11)

    add_heading(doc, '4. Elimina la fricción de los ponentes internacionales', level=2)
    p = doc.add_paragraph()
    r1 = p.add_run('Los ponentes no necesitan cuenta UNEMI, no suben archivos, no graban '
                   'localmente. Se unen al link como a cualquier Meet. El sistema hace '
                   'todo por detrás. ')
    r1.font.size = Pt(11)
    r2 = p.add_run('Invitar a un ponente de Stanford es igual de fácil que invitar al de al lado.')
    r2.bold = True
    r2.font.size = Pt(11)

    add_heading(doc, '5. Convierte a la IA en infraestructura, no en gimmick', level=2)
    p = doc.add_paragraph()
    r1 = p.add_run('No hay "asistente IA que responde en el chat". No hay "predicción de deserción". '
                   'La IA aquí hace ')
    r1.font.size = Pt(11)
    r2 = p.add_run('trabajo invisible')
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(': lee, entiende, sintetiza, publica. Es un editor silencioso '
                   'que convierte 2 horas de video en contenido consumible en 5 minutos.')
    r3.font.size = Pt(11)
    add_divider(doc)

    # SECCIÓN 5 — Valor en números
    add_heading(doc, 'El valor en números')
    add_body(doc, 'Imagina una facultad que organiza 30 eventos al año:')

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    headers = ['Métrica', 'Hoy', 'Con AI Minutes']
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)

    rows = [
        ('Tiempo manual post-evento', '3-4h por evento', '0 minutos'),
        ('Contenido reutilizable generado', '~0', '~30 resúmenes + 300 Q&As + 250 highlights'),
        ('Alcance post-evento', '0 personas/mes', 'Indefinido (SEO + compartibles)'),
        ('Costo por evento procesado', 'N/A', '~$0.05 en IA'),
        ('Trabajo humano para mantenerlo', 'Inviable', '1 clic'),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = table.rows[i].cells[j]
            c.text = val
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Un año de eventos se convierte en una biblioteca navegable permanente.')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = PRIMARY
    add_body(doc, 'Sin contratar a nadie nuevo. Sin pedir nada extra a los docentes.')
    add_divider(doc)

    # SECCIÓN 6 — Elegancia técnica
    add_heading(doc, 'Lo que lo hace técnicamente elegante')
    add_body(
        doc,
        'La mayoría de "integraciones con Zoom/Meet" que existen en el mercado son pesadas:'
    )
    add_bullet(doc, 'Bots que se unen a tu reunión (incómodo visualmente, límites de plan).')
    add_bullet(doc, 'SDKs propietarios costosos.')
    add_bullet(doc, 'Grabaciones que hay que subir manualmente.')
    p = doc.add_paragraph()
    r1 = p.add_run('Esta arquitectura ')
    r1.font.size = Pt(11)
    r2 = p.add_run('evita todo eso')
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(':')
    r3.font.size = Pt(11)
    add_bullet(doc, ' (transcripción nativa de Workspace Education).',
               bold_first='Usa lo que Google YA hace gratis')
    add_bullet(doc, ' — tu propia cuenta es el organizador, todo queda en tu Drive.',
               bold_first='No necesita bots externos')
    add_bullet(doc, ' — el video vive en Drive, solo se referencia.',
               bold_first='No duplica storage')
    add_bullet(doc, ' — las cuotas gratuitas de Google API cubren miles de eventos.',
               bold_first='Escalable por default')
    p = doc.add_paragraph()
    r1 = p.add_run('El costo marginal por evento procesado es ')
    r1.font.size = Pt(11)
    r2 = p.add_run('$0.05 en Claude API')
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = ACCENT
    r3 = p.add_run('. Nada más.')
    r3.font.size = Pt(11)
    add_divider(doc)

    # SECCIÓN 7 — Escenarios
    add_heading(doc, 'Escenarios que desbloquea')

    add_heading(doc, 'Para el estudiante', level=3)
    add_quote(
        doc,
        '"Perdí la charla del Dr. Martínez el martes. Abro Certifai, busco su nombre, '
        'leo el resumen en 3 minutos, veo las 10 preguntas que generó la IA. En 10 '
        'minutos tengo el 80% del valor de haber estado 2 horas."'
    )

    add_heading(doc, 'Para el docente', level=3)
    add_quote(
        doc,
        '"Di una clase magistral sobre ética de IA. Al día siguiente el sistema '
        'publica mi resumen automático. Lo comparto en LinkedIn. Me invitan a dar '
        'la misma charla en otra universidad."'
    )

    add_heading(doc, 'Para el director académico', level=3)
    add_quote(
        doc,
        '"¿Qué temas fueron más discutidos este semestre? El sistema me da una nube '
        'de temas clusterizados con la IA de todos los eventos. Puedo detectar '
        'tendencias, gaps curriculares, oportunidades."'
    )

    add_heading(doc, 'Para el estudiante externo que llegó por Google', level=3)
    add_quote(
        doc,
        '"Busqué \'modelos de lenguaje aplicados a educación\' y encontré el evento '
        'del Dr. López en Certifai. No sabía que existía UNEMI. Me inscribí al '
        'siguiente evento."'
    )
    add_divider(doc)

    # SECCIÓN 8 — Tesis
    add_heading(doc, 'Por qué esto es una tesis, no solo un producto')
    add_body(doc, 'Como trabajo académico, esta solución:')
    add_bullet(doc, ' (Google Workspace, Claude API, Django). Demuestra dominio de arquitecturas distribuidas.',
               bold_first='Integra 3 sistemas heterogéneos')
    add_bullet(doc, ': prompt engineering, prompt caching, structured output, multi-step reasoning.',
               bold_first='Usa técnicas modernas de IA aplicada')
    add_bullet(doc, ': precisión del resumen vs humano, tiempo ahorrado, retención de contenido.',
               bold_first='Es evaluable cuantitativamente')
    add_bullet(doc, ': la generación automática de Q&A se apoya en la taxonomía de Bloom (comprensión, aplicación, análisis).',
               bold_first='Tiene fundamento pedagógico')
    add_bullet(doc, ': cualquier universidad con Google Workspace puede adoptarlo en semanas.',
               bold_first='Es replicable')
    add_bullet(doc, ': no existe un sistema open-source equivalente para educación continua universitaria en español.',
               bold_first='Genera contribución original')
    add_divider(doc)

    # SECCIÓN 9 — 2 semanas
    add_heading(doc, 'Lo que se logra en 2 semanas de trabajo')
    p = doc.add_paragraph()
    r1 = p.add_run('Con la arquitectura ya diseñada, el MVP está a ')
    r1.font.size = Pt(11)
    r2 = p.add_run('10-14 días')
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(' de trabajo enfocado:')
    r3.font.size = Pt(11)
    add_bullet(doc, ' OAuth con Google + Calendar API + descarga de transcripts + pipeline Claude.',
               bold_first='Semana 1:')
    add_bullet(doc, ' Detección automática + UI pública del resumen + testing end-to-end.',
               bold_first='Semana 2:')
    p = doc.add_paragraph()
    r1 = p.add_run('Todo dentro del cronograma de tesis. ')
    r1.font.size = Pt(11)
    r2 = p.add_run('Es factible y ambicioso a la vez')
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(' — la combinación exacta que un tribunal respeta.')
    r3.font.size = Pt(11)
    add_divider(doc)

    # SECCIÓN 10 — La apuesta
    add_heading(doc, 'La apuesta')
    p = doc.add_paragraph()
    r = p.add_run(
        'Construimos esto y Certifai deja de ser "una plataforma de certificados" '
        'para convertirse en algo que no tiene nombre todavía:'
    )
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = PRIMARY
    add_bullet(doc, 'No es LMS.')
    add_bullet(doc, 'No es herramienta de videoconferencia.')
    add_bullet(doc, 'No es un generador de notas con IA.')
    p = doc.add_paragraph()
    r1 = p.add_run('Es el ')
    r1.font.size = Pt(11)
    r2 = p.add_run('cerebro compartido')
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = ACCENT
    r3 = p.add_run(' de una universidad. La capa que captura cada conversación educativa '
                   'y la vuelve accesible, navegable, perpetua.')
    r3.font.size = Pt(11)
    add_body(doc, 'Cada evento que organiza UNEMI lo hace una institución más inteligente.',
             italic=True)
    add_body(doc, 'Cada evento que se pierde es contenido que no va a existir nunca.',
             italic=True)
    add_divider(doc)

    # SECCIÓN 11 — Una última cosa
    add_heading(doc, 'Una última cosa')
    add_body(doc, 'Si esto te convence, el siguiente paso es pequeño:')
    add_bullet(doc, 'Conectar la cuenta UNEMI al proyecto (una vez).')
    add_bullet(doc, 'Crear una sesión virtual de prueba.')
    add_bullet(doc, 'Hacer una reunión de 15 minutos de cualquier cosa.')
    add_bullet(doc, 'Ver cómo, automáticamente, en tu página aparece un resumen con timeline y Q&A.')
    add_body(
        doc,
        'Ese momento — cuando veas tu primera reunión convertida en contenido sin '
        'haber tocado un solo archivo — es cuando entiendes que estamos construyendo '
        'algo diferente.',
        italic=True, color=GRAY,
    )
    add_divider(doc)

    # Cierre
    add_title(doc, 'Certifai AI Minutes', size=20)
    add_body(
        doc,
        'Porque el conocimiento no debería morir cuando termina la reunión.',
        italic=True, color=ACCENT, size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    out = Path(__file__).parent / 'Certifai_AI_Minutes_Pitch.docx'
    doc.save(out)
    print(f'OK: {out}')


if __name__ == '__main__':
    build()
