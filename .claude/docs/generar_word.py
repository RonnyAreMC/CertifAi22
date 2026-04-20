"""Genera el documento Word del analisis de caso - CertifAI"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BLUE = RGBColor(0x16, 0x20, 0x54)
ORANGE = RGBColor(0xF5, 0x88, 0x30)
GRAY = RGBColor(0x64, 0x74, 0x8B)
BLACK = RGBColor(0x1E, 0x29, 0x3B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11.5)
style.font.color.rgb = BLACK
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)


def blank(n=1):
    for _ in range(n):
        doc.add_paragraph()

def colored_line():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 60)
    run.font.color.rgb = ORANGE
    run.font.size = Pt(6)

def heading_num(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    run_num = p.add_run(f'  {num}  ')
    run_num.font.size = Pt(11)
    run_num.font.bold = True
    run_num.font.color.rgb = WHITE
    shd = run_num._element.get_or_add_rPr()
    bg = shd.makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):'F58830'})
    shd.append(bg)
    run_text = p.add_run(f'   {text}')
    run_text.font.size = Pt(14)
    run_text.font.bold = True
    run_text.font.color.rgb = BLUE
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'):'single',qn('w:sz'):'8',qn('w:space'):'4',qn('w:color'):'162054'})
    pBdr.append(bottom)
    pPr.append(pBdr)

def body(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.color.rgb = BLACK
    run.font.bold = bold

def bullet(text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = BLACK
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = BLACK

def highlight_box(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    left = pBdr.makeelement(qn('w:left'), {qn('w:val'):'single',qn('w:sz'):'24',qn('w:space'):'8',qn('w:color'):'F58830'})
    pBdr.append(left)
    pPr.append(pBdr)
    shd = pPr.makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):'FEF3E8'})
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.color.rgb = BLUE
    run.font.italic = True

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = WHITE
        shd = cell._element.get_or_add_tcPr()
        bg = shd.makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):'162054'})
        shd.append(bg)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK
            if r_idx % 2 == 1:
                shd = cell._element.get_or_add_tcPr()
                bg = shd.makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):'E8ECF4'})
                shd.append(bg)


# ═══════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════
blank(4)
colored_line()
blank(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ANALISIS DE CASO')
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = ORANGE
run.font.letter_spacing = Pt(3)

blank(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CertifAI: plataforma web con inteligencia artificial para la emision y verificacion de certificados academicos')
run.font.size = Pt(17)
run.font.bold = True
run.font.color.rgb = BLUE

blank(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Propuesta de automatizacion de procesos administrativos en instituciones de educacion superior')
run.font.size = Pt(10.5)
run.font.color.rgb = GRAY
run.font.italic = True

blank(2)
colored_line()
blank(1)

# Grupo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Grupo 22')
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = BLUE

blank(0)

integrantes = [
    'Arellano Urgiles Ronny Isaac',
    'Dabor Enrique Garcia Sanchez',
]
for nombre in integrantes:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nombre)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

blank(1)

meta = [
    ('Modalidad:', ' Proyecto de titulacion de grado'),
    ('Entrega:', ' Primera entrega'),
    ('Fecha:', ' Abril 2026'),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label)
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = BLUE
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY

blank(1)

# Link arbol
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Arbol de problemas interactivo: ')
run.font.size = Pt(9.5)
run.font.color.rgb = GRAY
from docx.oxml import OxmlElement
hyperlink = OxmlElement('w:hyperlink')
hyperlink.set(qn('w:history'), '1')
r_id = doc.part.relate_to('https://certifaiprot.netlify.app/', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
hyperlink.set(qn('r:id'), r_id)
new_run = OxmlElement('w:r')
rPr = OxmlElement('w:rPr')
rStyle = OxmlElement('w:rStyle')
rStyle.set(qn('w:val'), 'Hyperlink')
rPr.append(rStyle)
sz = OxmlElement('w:sz')
sz.set(qn('w:val'), '19')
rPr.append(sz)
color_el = OxmlElement('w:color')
color_el.set(qn('w:val'), 'F58830')
rPr.append(color_el)
new_run.append(rPr)
new_run.text = 'https://certifaiprot.netlify.app/'
hyperlink.append(new_run)
p._element.append(hyperlink)

doc.add_page_break()


# ═══════════════════════════════════════════════
# 1. TITULO
# ═══════════════════════════════════════════════
heading_num('1', 'Titulo del analisis de caso')
body('CertifAI: plataforma web con inteligencia artificial para la emision y verificacion de certificados academicos.', bold=True)


# ═══════════════════════════════════════════════
# 2. INTRODUCCION
# ═══════════════════════════════════════════════
heading_num('2', 'Introduccion')

body('Las instituciones de educacion superior organizan anualmente seminarios, talleres y capacitaciones que requieren la emision de certificados como constancia de participacion o aprobacion. Este proceso, aparentemente sencillo, involucra una cadena de actividades que en la practica se realiza de forma manual y fragmentada.')

body('El flujo tipico comienza con la recopilacion de datos de asistencia mediante formularios web externos. Las respuestas se exportan como hojas de calculo que deben limpiarse manualmente: eliminar registros duplicados, corregir variaciones en nombres, validar formatos de identificacion y normalizar correos electronicos. Una vez depurados, los datos se cargan en una plataforma diferente para generar los certificados en formato PDF.')

body('La distribucion de los documentos presenta dos caminos igualmente problematicos. En el caso de certificados impresos, cada documento debe llevarse fisicamente a las autoridades para obtener su firma manuscrita. En el caso digital, cada archivo se adjunta y envia individualmente por correo electronico al participante.')

body('Este flujo manual genera consecuencias directas: errores en los datos de los certificados (como mostrar el correo del participante donde deberia aparecer el nombre del evento), duplicacion de registros de participantes, retrasos de semanas en la entrega, e imposibilidad de verificar la autenticidad de los documentos emitidos. Estos problemas se agravan proporcionalmente con el crecimiento en el numero de eventos y participantes.')

body('El presente trabajo propone el desarrollo de un prototipo que centralice este proceso en una unica plataforma web, incorporando inteligencia artificial para automatizar las tareas mas criticas y propensas a error humano.')


# ═══════════════════════════════════════════════
# 3. OBJETIVO
# ═══════════════════════════════════════════════
heading_num('3', 'Objetivo general')

highlight_box('Desarrollar un prototipo de sistema web que automatice la gestion de certificados academicos en instituciones de educacion superior, integrando modulos de inteligencia artificial para la clasificacion automatica de datos, generacion asistida de contenido, recomendaciones personalizadas e interfaces conversacionales, como alternativa a los procesos manuales fragmentados actuales.')


# ═══════════════════════════════════════════════
# 4. JUSTIFICACION
# ═══════════════════════════════════════════════
heading_num('4', 'Justificacion')

body('La digitalizacion de procesos administrativos en el ambito educativo es una necesidad reconocida internacionalmente. Sin embargo, muchas instituciones de educacion superior, particularmente las publicas en Latinoamerica, mantienen flujos operativos manuales para la gestion de certificados que generan problemas recurrentes y documentables.')

body('Los problemas tipicos identificados durante la fase de analisis incluyen:')

bullet(' al asociar manualmente las columnas de una hoja de calculo con los campos del certificado, es frecuente asignar la columna equivocada. Esto produce certificados con datos cruzados, como el correo del participante en el espacio destinado al nombre del evento.', bold_prefix='Errores por mapeo incorrecto de datos:')
bullet(' sin una entidad centralizada de participante con logica de cruce por identificacion o correo, la misma persona puede aparecer multiples veces con variaciones menores en sus datos.', bold_prefix='Duplicacion de registros:')
bullet(' los certificados generados no incluyen mecanismos de autenticidad (codigos QR, hash criptografico), lo que los hace susceptibles a falsificacion.', bold_prefix='Ausencia de verificacion:')
bullet(' la necesidad de obtener firmas fisicas o enviar certificados uno a uno por correo genera retrasos que pueden extenderse por semanas.', bold_prefix='Cuellos de botella en firmas y distribucion:')

body('La integracion de inteligencia artificial no se plantea como un componente decorativo sino como respuesta directa a estos problemas: un modelo de lenguaje puede clasificar columnas de datos por su contenido real (no por su encabezado), asistir en la redaccion de textos institucionales, generar recomendaciones personalizadas basadas en historial, y facilitar la creacion de eventos mediante comandos de voz con extraccion de entidades.')


# ═══════════════════════════════════════════════
# 5. SOLUCION PROPUESTA
# ═══════════════════════════════════════════════
heading_num('5', 'Solucion propuesta')

body('Se propone un prototipo de sistema web construido con Django (Python) que centralice el ciclo completo de gestion de certificados. El prototipo integra cinco modulos de inteligencia artificial sobre un nucleo funcional de gestion de certificados y control de asistencia:')

bullet(' generacion dinamica de documentos PDF con plantillas personalizables, firmas digitales precargadas y codigos QR de verificacion unica.', bold_prefix='Gestion de certificados:')
bullet(' registro mediante codigos QR, confirmacion de cupos y generacion automatica de lotes de certificados a partir de sesiones.', bold_prefix='Control de asistencia:')
bullet(' un modelo de lenguaje analiza el contenido de las columnas de la hoja de calculo (no solo sus encabezados) para determinar automaticamente que campo representa cada columna.', bold_prefix='IA para clasificacion de datos:')
bullet(' generacion asistida del texto del certificado, con conocimiento de las variables del sistema y el contexto academico.', bold_prefix='IA como asistente de redaccion:')
bullet(' sistema basado en embeddings vectoriales que sugiere eventos futuros relevantes a cada participante segun su historial de asistencia.', bold_prefix='IA para recomendaciones:')
bullet(' interfaz que permite crear eventos dictando por voz, con extraccion automatica de entidades (fecha, hora, capacidad, titulo).', bold_prefix='IA para comandos de voz:')
bullet(' generacion de narrativa ejecutiva a partir de las metricas del sistema para facilitar la toma de decisiones.', bold_prefix='IA para insights:')

body('Adicionalmente, el prototipo expone una API REST documentada con autenticacion JWT, preparada para la futura integracion con aplicaciones moviles u otros sistemas.')


# ═══════════════════════════════════════════════
# 6. CASOS SIMILARES
# ═══════════════════════════════════════════════
doc.add_page_break()
heading_num('6', 'Casos similares e investigaciones previas')

body('Se identificaron los siguientes antecedentes relevantes, tanto comerciales como academicos:')

add_table(
    ['Referencia', 'Tipo', 'Descripcion', 'Diferencia con CertifAI'],
    [
        ['Accredible (EE.UU., 2013)', 'Plataforma comercial', 'Credenciales digitales verificables con blockchain', 'Sin IA para clasificacion ni recomendaciones. Costo inaccesible para universidades publicas'],
        ['Certifier.io (Polonia, 2019)', 'Plataforma comercial', 'Generacion masiva de certificados desde CSV/Excel', 'Sin asistencia, sin verificacion QR, sin asistente de IA'],
        ['Automated Certificate Generation (IEEE, 2022)', 'Articulo cientifico', 'Python y ReportLab para generacion automatizada', 'No contempla IA ni verificacion de autenticidad'],
        ['Recommendation Systems in Education (Comp. & Ed., 2021)', 'Revision sistematica', 'Content-based filtering con embeddings supera metodos tradicionales', 'Valida el enfoque del modulo de recomendaciones'],
        ['Voice-controlled Educational Interfaces (ACM, 2023)', 'Articulo cientifico', 'Entrada por voz reduce tiempo de ingreso en 40%', 'Valida la interfaz de voz propuesta'],
        ['LLMs for Tabular Data Classification (NeurIPS, 2023)', 'Articulo cientifico', 'LLMs clasifican datos tabulares con >90% precision', 'Sustenta el modulo de clasificacion inteligente'],
    ]
)

blank(1)
body('Trabajos de grado relacionados en la region:', bold=True)

add_table(
    ['Trabajo', 'Institucion', 'Anio', 'Diferencia con CertifAI'],
    [
        ['Sistema web para gestion de certificados digitales', 'ESPOCH', '2021', 'Solo generacion basica sin IA ni verificacion'],
        ['Plataforma con firma electronica', 'UTA', '2022', 'Enfocada solo en firma, sin gestion integral'],
        ['Recomendacion de cursos con Machine Learning', 'UTPL', '2023', 'Solo recomendaciones, sin certificados ni asistencia'],
        ['Sistema de asistencia con codigos QR', 'UG', '2022', 'Solo asistencia QR, sin certificados ni IA'],
    ]
)

blank(1)
highlight_box('Diferenciador de CertifAI: ninguno de los antecedentes identificados integra en un solo prototipo la gestion de certificados con multiples plantillas, control de asistencia QR, cinco modulos de inteligencia artificial, deduplicacion de participantes y API REST. Esta integracion constituye la contribucion principal de la propuesta.')

# Footer
blank(2)
colored_line()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Analisis de caso  |  Grupo 22  |  Primera entrega  |  Abril 2026')
run.font.size = Pt(8)
run.font.color.rgb = GRAY

# Save
output = os.path.join(os.path.dirname(__file__), 'CertifAI_Analisis_de_Caso.docx')
doc.save(output)
print(f'Word generado: {output}')
